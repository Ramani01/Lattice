"""Remove literal asterisks from all editable text in the research-paper DOCX."""
from pathlib import Path
from docx import Document

PATH = Path(__file__).parent / "Lattice_Research_Paper_Draft.docx"


def clean_paragraphs(paragraphs):
    changed = 0
    for paragraph in paragraphs:
        for run in paragraph.runs:
            if "*" in run.text:
                changed += run.text.count("*")
                run.text = run.text.replace("*", "")
    return changed


def clean_table(table):
    changed = 0
    for row in table.rows:
        for cell in row.cells:
            changed += clean_paragraphs(cell.paragraphs)
            for nested in cell.tables:
                changed += clean_table(nested)
    return changed


document = Document(PATH)
removed = clean_paragraphs(document.paragraphs)
for table in document.tables:
    removed += clean_table(table)
for section in document.sections:
    removed += clean_paragraphs(section.header.paragraphs)
    removed += clean_paragraphs(section.footer.paragraphs)
    for table in section.header.tables:
        removed += clean_table(table)
    for table in section.footer.tables:
        removed += clean_table(table)
document.save(PATH)
print(f"Removed {removed} asterisk character(s) from {PATH.name}.")
