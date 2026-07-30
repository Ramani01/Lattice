"""Build a polished Word version of the Lattice research-paper draft."""
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).parent
SOURCE = ROOT / "paper.md"
OUTPUT = ROOT / "Lattice_Research_Paper_Draft.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_FILL = "F4F6F9"


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    elem = OxmlElement("w:shd")
    elem.set(qn("w:fill"), fill)
    props.append(elem)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure(doc):
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "LATTICE | RESEARCH PAPER DRAFT"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.runs[0].font.name = "Calibri"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor.from_string("6B7280")
    footer = section.footer.paragraphs[0]
    footer_run = footer.add_run()
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor.from_string("6B7280")
    add_page_number(footer)


def add_table(doc, rows):
    cols = len(rows[0])
    table = doc.add_table(rows=0, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    usable = 6.5
    widths = [usable / cols] * cols
    if cols == 2:
        widths = [1.85, 4.65]
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                shade(cell, LIGHT_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            run = p.add_run(value.replace("**", ""))
            run.font.size = Pt(9)
            if r_idx == 0:
                run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def is_table_line(line):
    return line.startswith("|") and line.endswith("|")


def parse_table_line(line):
    return [part.strip() for part in line.strip()[1:-1].split("|")]


def add_markdown(doc, source):
    lines = source.splitlines()
    i = 0
    in_code = False
    code_lines = []
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(8)
                r = p.add_run("\n".join(code_lines))
                r.font.name = "Consolas"
                r.font.size = Pt(8.5)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if is_table_line(line):
            rows = []
            while i < len(lines) and is_table_line(lines[i].rstrip()):
                candidate = parse_table_line(lines[i].rstrip())
                if not all(re.fullmatch(r":?-{2,}:?", item.replace(" ", "")) for item in candidate):
                    rows.append(candidate)
                i += 1
            if rows:
                add_table(doc, rows)
            continue
        if line.startswith("# "):
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif re.match(r"^\d+\. ", line):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip() in {"---", ""}:
            if line == "":
                pass
        else:
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Remove Markdown emphasis markers for a clean, readable draft.
            p.add_run(line.replace("**", "").replace("*", ""))
        i += 1


def build():
    doc = Document()
    configure(doc)

    # Editorial-cover first page (named override: centered academic cover).
    for _ in range(5):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("LATTICA: GRAPH-RAG-BASED INFRASTRUCTURE IMPACT ANALYSIS AND SAFE EXECUTION PLANNING FRAMEWORK")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    title.paragraph_format.space_after = Pt(16)
    subtitle = doc.add_paragraph("Capstone Research Paper Draft")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor.from_string("4B5563")
    for _ in range(8):
        doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Prepared by: [Student Name]\n").bold = True
    meta.add_run("[Department / College]\n[Submission Date]")
    doc.add_page_break()

    add_markdown(doc, SOURCE.read_text(encoding="utf-8"))
    doc.core_properties.title = "Lattica Research Paper Draft"
    doc.core_properties.subject = "Microservice migration planning"
    doc.core_properties.author = "[Student Name]"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
